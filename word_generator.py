# -*- coding: utf-8 -*-
"""
word_generator.py
------------------------------------------------------------------
Modul untuk menghasilkan dokumen laporan keuangan resmi dalam
format Microsoft Word (.docx) menggunakan python-docx.

Mencakup: kop surat + logo, judul, nomor surat, paragraf pembuka,
tabel transaksi, ringkasan keuangan, grafik, paragraf penutup,
blok tanda tangan, watermark, header/footer otomatis, nomor
halaman, dan QR Code verifikasi.
------------------------------------------------------------------
"""

import os
import sys
import tempfile
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from chart_utils import buat_grafik_pemasukan_pengeluaran, buat_qrcode

HIJAU_TUA = RGBColor(0x1B, 0x5E, 0x20)
_BASE_DIR_WORD = os.path.dirname(sys.executable) if getattr(sys, "frozen", False) \
    else os.path.dirname(os.path.abspath(__file__))
ABU_MUDA = "F2F2F2"


class WordGenerator:
    """Bertanggung jawab membangun dokumen .docx laporan keuangan BUMDes."""

    def __init__(self, laporan, logo_path=None):
        """
        Parameters
        ----------
        laporan : LaporanKeuangan
            Objek laporan berisi data transaksi + ringkasan + info BUMDes.
        logo_path : str
            Path ke file logo BUMDes (opsional, pakai logo default jika kosong).
        """
        self.laporan = laporan
        self.info = laporan.info_bumdes
        self.logo_path = logo_path if logo_path and os.path.exists(logo_path) else \
            os.path.join(_BASE_DIR_WORD, "assets", "logo_default.png")
        self.doc = Document()
        self._temp_files = []

    # ------------------------------------------------------------------
    # ENTRY POINT
    # ------------------------------------------------------------------
    def generate(self, output_path):
        """Membangun seluruh isi dokumen lalu menyimpannya ke output_path."""
        self._setup_page()
        self._tambah_header_footer()
        self._tambah_watermark()
        self._tambah_kop_surat()
        self._tambah_judul()
        self._tambah_paragraf_pembuka()
        self._tambah_tabel_transaksi()
        self._tambah_ringkasan()
        self._tambah_grafik()
        self._tambah_paragraf_penutup()
        self._tambah_tanda_tangan()
        self._tambah_qrcode()

        self.doc.save(output_path)
        self._bersihkan_file_sementara()
        return output_path

    # ------------------------------------------------------------------
    # PENGATURAN HALAMAN
    # ------------------------------------------------------------------
    def _setup_page(self):
        section = self.doc.sections[0]
        # Landscape A4 (mendatar) — eksplisit A4 (bukan default Letter bawaan
        # template Word), supaya konsisten dengan ukuran kertas versi PDF.
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)

        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.6)

        style = self.doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

    # ------------------------------------------------------------------
    # HEADER / FOOTER OTOMATIS + NOMOR HALAMAN
    # ------------------------------------------------------------------
    def _tambah_header_footer(self):
        section = self.doc.sections[0]

        # --- Footer: nama BUMDes kiri, nomor halaman kanan ---
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = ""
        run = p.add_run(f"{self.info.get('nama_bumdes', 'BUMDes')} — Dokumen Resmi")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

        tab_stops = p.paragraph_format.tab_stops
        section_width = section.page_width - section.left_margin - section.right_margin
        tab_stops.add_tab_stop(section_width, alignment=2)  # rata kanan

        run_tab = p.add_run("\t")
        self._tambah_field_nomor_halaman(p)

    def _tambah_field_nomor_halaman(self, paragraph):
        """Menambahkan field otomatis 'Halaman X dari Y' ke paragraf footer."""
        run = paragraph.add_run("Halaman ")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

        self._insert_field(paragraph, "PAGE")
        run2 = paragraph.add_run(" dari ")
        run2.font.size = Pt(8)
        run2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        self._insert_field(paragraph, "NUMPAGES")

    @staticmethod
    def _insert_field(paragraph, field_code):
        """Menyisipkan field Word (mis. PAGE, NUMPAGES) secara manual via XML."""
        run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" {field_code} "

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)

    # ------------------------------------------------------------------
    # WATERMARK "DOKUMEN RESMI BUMDES"
    # ------------------------------------------------------------------
    def _tambah_watermark(self):
        """Menyisipkan watermark teks diagonal transparan di header setiap halaman."""
        section = self.doc.sections[0]
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        # Textbox berputar berisi teks watermark, ditempatkan di belakang teks.
        xml = f'''
        <w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                xmlns:v="urn:schemas-microsoft-com:vml"
                xmlns:o="urn:schemas-microsoft-com:office:office">
            <v:shapetype id="_x0000_t136" coordsize="1600,21600" o:spt="136"
                adj="10800" path="m@7,0l@8,0m@5,21600l@6,21600e">
            </v:shapetype>
            <v:shape id="WatermarkShape" type="#_x0000_t136"
                style="position:absolute;margin-left:0;margin-top:0;width:420pt;height:110pt;
                       z-index:-251654144;mso-position-horizontal:center;
                       mso-position-horizontal-relative:margin;
                       mso-position-vertical:center;
                       mso-position-vertical-relative:margin;rotation:315"
                o:allowincell="f" fillcolor="#1B5E20" stroked="f">
                <v:fill opacity="14000f"/>
                <v:textpath style="font-family:'Calibri';font-size:40pt;font-weight:bold"
                    string="DOKUMEN RESMI BUMDES"/>
            </v:shape>
        </w:pict>
        '''
        run._r.append(OxmlElement("w:pict")) if False else None
        # Gunakan pendekatan parsing XML langsung agar valid.
        from docx.oxml import parse_xml
        el = parse_xml(xml.strip())
        run._r.append(el)

    # ------------------------------------------------------------------
    # KOP SURAT
    # ------------------------------------------------------------------
    def _tambah_kop_surat(self):
        table = self.doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Cm(3.2)
        table.columns[1].width = Cm(23.3)

        # Kolom logo
        cell_logo = table.cell(0, 0)
        cell_logo.vertical_alignment = 1
        p_logo = cell_logo.paragraphs[0]
        run_logo = p_logo.add_run()
        try:
            run_logo.add_picture(self.logo_path, width=Cm(2.6))
        except Exception:
            p_logo.add_run("[LOGO]")

        # Kolom informasi BUMDes
        cell_info = table.cell(0, 1)
        p1 = cell_info.paragraphs[0]
        r1 = p1.add_run("BADAN USAHA MILIK DESA (BUMDes)")
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = HIJAU_TUA

        p2 = cell_info.add_paragraph()
        r2 = p2.add_run(self.info.get("nama_bumdes", "-"))
        r2.bold = True
        r2.font.size = Pt(15)
        r2.font.color.rgb = HIJAU_TUA

        alamat_line = (
            f"Desa {self.info.get('nama_desa', '-')}, "
            f"Kec. {self.info.get('kecamatan', '-')}, "
            f"Kab. {self.info.get('kabupaten', '-')}"
        )
        p3 = cell_info.add_paragraph()
        r3 = p3.add_run(alamat_line)
        r3.font.size = Pt(9)

        kontak_bagian = []
        if self.info.get("telepon"):
            kontak_bagian.append(f"Telp: {self.info.get('telepon')}")
        if self.info.get("email"):
            kontak_bagian.append(f"Email: {self.info.get('email')}")
        if self.info.get("website"):
            kontak_bagian.append(f"Website: {self.info.get('website')}")
        if kontak_bagian:
            p4 = cell_info.add_paragraph()
            r4 = p4.add_run(" | ".join(kontak_bagian))
            r4.font.size = Pt(9)

        # Hilangkan border tabel kop surat
        self._hilangkan_border_tabel(table)

        # Garis pembatas
        self._tambah_garis_horizontal()
        self.doc.add_paragraph()

    def _tambah_garis_horizontal(self):
        p = self.doc.add_paragraph()
        p_format = p.paragraph_format
        p_format.space_before = Pt(2)
        p_format.space_after = Pt(2)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "18")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1B5E20")
        pBdr.append(bottom)
        pPr.append(pBdr)

    @staticmethod
    def _hilangkan_border_tabel(table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "nil")
            borders.append(el)
        tblPr.append(borders)

    # ------------------------------------------------------------------
    # JUDUL + META
    # ------------------------------------------------------------------
    def _tambah_judul(self):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("LAPORAN KEUANGAN")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = HIJAU_TUA

        meta = [
            ("Nomor Surat", self.info.get("nomor_surat", "-")),
            ("Periode", self.info.get("periode", "-")),
            ("Tanggal Cetak", self.laporan.tanggal_cetak()),
        ]
        table = self.doc.add_table(rows=len(meta), cols=3)
        table.autofit = False
        table.columns[0].width = Cm(3.2)
        table.columns[1].width = Cm(0.4)
        table.columns[2].width = Cm(9)
        for i, (label, value) in enumerate(meta):
            table.cell(i, 0).paragraphs[0].add_run(label).font.size = Pt(10)
            table.cell(i, 1).paragraphs[0].add_run(":").font.size = Pt(10)
            run_val = table.cell(i, 2).paragraphs[0].add_run(str(value))
            run_val.font.size = Pt(10)
            run_val.bold = True
        self._hilangkan_border_tabel(table)
        self.doc.add_paragraph()

    def _tambah_paragraf_pembuka(self):
        p = self.doc.add_paragraph(self.laporan.paragraf_pembuka())
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(10)

    # ------------------------------------------------------------------
    # TABEL TRANSAKSI
    # ------------------------------------------------------------------
    # Lebar kolom (cm), total ~26 cm — memanfaatkan lebar penuh halaman
    # landscape (section width sekitar 26.5cm), supaya mayoritas baris
    # transaksi muat dalam SATU baris saja (mirip tampilan Excel).
    LEBAR_KOLOM_TRANSAKSI = [1.0, 2.1, 2.3, 7.0, 2.1, 2.6, 2.6, 2.6, 3.2]

    def _tambah_tabel_transaksi(self):
        headers = ["No", "Tanggal", "No. Transaksi", "Uraian", "Kategori",
                   "Pemasukan", "Pengeluaran", "Saldo", "Keterangan"]
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._set_tabel_layout_fixed(table)

        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].width = Cm(self.LEBAR_KOLOM_TRANSAKSI[i])
            hdr_cells[i].text = h
            for p in hdr_cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            self._set_cell_background(hdr_cells[i], "1B5E20")

        for idx, t in enumerate(self.laporan.transaksi_list, start=1):
            row = table.add_row().cells
            nilai = [
                str(idx),
                self.laporan.format_tanggal(t["tanggal"]),
                t["nomor_transaksi"],
                t["uraian"],
                t["kategori"],
                self.laporan.format_rupiah(t["pemasukan"]) if t["pemasukan"] else "-",
                self.laporan.format_rupiah(t["pengeluaran"]) if t["pengeluaran"] else "-",
                self.laporan.format_rupiah(t["saldo"]),
                t["keterangan"],
            ]
            for i, val in enumerate(nilai):
                row[i].width = Cm(self.LEBAR_KOLOM_TRANSAKSI[i])
                row[i].text = val
                for p in row[i].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 1, 2) else WD_ALIGN_PARAGRAPH.LEFT
                    for r in p.runs:
                        r.font.size = Pt(9)
            if idx % 2 == 0:
                for cell in row:
                    self._set_cell_background(cell, ABU_MUDA)

        # Terapkan ulang lebar kolom di level tabel (python-docx butuh ini
        # dipanggil juga di sini agar Word konsisten menampilkan lebar yang sama
        # di semua baris, termasuk baris yang ditambahkan lewat add_row()).
        for col, lebar in zip(table.columns, self.LEBAR_KOLOM_TRANSAKSI):
            col.width = Cm(lebar)

        self.doc.add_paragraph()

    @staticmethod
    def _set_tabel_layout_fixed(table):
        """
        Memaksa Word memakai lebar kolom tetap (fixed layout) alih-alih
        menebak sendiri lebar tiap kolom berdasarkan isi teks — ini yang
        menyebabkan kolom "Uraian" tabrakan dengan kolom sebelahnya kalau
        dibiarkan default (autofit to contents).
        """
        tbl = table._tbl
        tblPr = tbl.tblPr
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tblPr.append(layout)

    @staticmethod
    def _set_cell_background(cell, hex_color):
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        cell._tc.get_or_add_tcPr().append(shd)

    # ------------------------------------------------------------------
    # RINGKASAN KEUANGAN
    # ------------------------------------------------------------------
    def _tambah_ringkasan(self):
        p = self.doc.add_paragraph()
        run = p.add_run("Ringkasan Keuangan")
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = HIJAU_TUA

        r = self.laporan.get_ringkasan_format()

        # --- 3 kartu statistik utama (Total Pemasukan/Pengeluaran/Saldo) ---
        kartu = self.doc.add_table(rows=2, cols=3)
        kartu.autofit = False
        self._set_tabel_layout_fixed(kartu)
        warna_kartu = ["1B5E20", "B71C1C", "37474F"]
        label_kartu = ["TOTAL PEMASUKAN", "TOTAL PENGELUARAN", "SALDO AKHIR"]
        nilai_kartu = [r["total_pemasukan"], r["total_pengeluaran"], r["saldo_akhir"]]

        for i in range(3):
            kartu.cell(0, i).width = Cm(8.6)
            kartu.cell(1, i).width = Cm(8.6)
            self._set_cell_background(kartu.cell(0, i), warna_kartu[i])
            self._set_cell_background(kartu.cell(1, i), warna_kartu[i])

            p_label = kartu.cell(0, i).paragraphs[0]
            p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_label = p_label.add_run(label_kartu[i])
            r_label.bold = True
            r_label.font.size = Pt(9)
            r_label.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            p_nilai = kartu.cell(1, i).paragraphs[0]
            p_nilai.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_nilai = p_nilai.add_run(nilai_kartu[i])
            r_nilai.bold = True
            r_nilai.font.size = Pt(15)
            r_nilai.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        self._hilangkan_border_tabel(kartu)
        self.doc.add_paragraph()

        # --- Detail tambahan (jumlah transaksi, rata-rata, terbesar/terkecil) ---
        data = [
            ("Jumlah Transaksi", r["jumlah_transaksi"]),
            ("Rata-rata Pemasukan", r["rata_pemasukan"]),
            ("Rata-rata Pengeluaran", r["rata_pengeluaran"]),
            ("Transaksi Terbesar", r["transaksi_terbesar"]),
            ("Transaksi Terkecil", r["transaksi_terkecil"]),
        ]
        table = self.doc.add_table(rows=len(data), cols=2)
        table.style = "Table Grid"
        table.autofit = False
        self._set_tabel_layout_fixed(table)
        for i, (label, value) in enumerate(data):
            table.cell(i, 0).width = Cm(6)
            table.cell(i, 1).width = Cm(20)
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)
            table.cell(i, 0).paragraphs[0].runs[0].font.size = Pt(10)
            table.cell(i, 0).paragraphs[0].runs[0].bold = True
            table.cell(i, 1).paragraphs[0].runs[0].font.size = Pt(10)
            self._set_cell_background(table.cell(i, 0), "E8F5E9")
        self.doc.add_paragraph()

    # ------------------------------------------------------------------
    # GRAFIK
    # ------------------------------------------------------------------
    def _tambah_grafik(self):
        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        tmp.close()
        self._temp_files.append(tmp.name)
        buat_grafik_pemasukan_pengeluaran(
            self.laporan.ringkasan["total_pemasukan"],
            self.laporan.ringkasan["total_pengeluaran"],
            tmp.name,
        )
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(tmp.name, width=Cm(15))
        self.doc.add_paragraph()

    def _tambah_paragraf_penutup(self):
        p = self.doc.add_paragraph(self.laporan.paragraf_penutup())
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(14)

    # ------------------------------------------------------------------
    # TANDA TANGAN
    # ------------------------------------------------------------------
    def _tambah_tanda_tangan(self):
        p_tempat = self.doc.add_paragraph()
        p_tempat.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_tempat.add_run(
            f"{self.info.get('nama_desa', '-')}, {self.laporan.tanggal_cetak()}"
        ).font.size = Pt(10)

        table = self.doc.add_table(rows=1, cols=2)
        table.autofit = True
        c1, c2 = table.rows[0].cells

        for cell, jabatan, nama in [
            (c1, "Direktur BUMDes", self.info.get("nama_direktur", "-")),
            (c2, "Bendahara", self.info.get("nama_bendahara", "-")),
        ]:
            p1 = cell.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p1.add_run(jabatan).font.size = Pt(10)

            for _ in range(4):
                cell.add_paragraph()

            # Garis pendek di atas nama (bukan underline pada teks nama)
            p_garis = cell.add_paragraph()
            p_garis.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._tambah_garis_pendek(p_garis)

            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(nama)
            r2.bold = True
            r2.font.size = Pt(10)

        self._hilangkan_border_tabel(table)

    @staticmethod
    def _tambah_garis_pendek(paragraph):
        """Menambahkan garis horizontal pendek di bawah paragraf (untuk tempat tanda tangan)."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)
        # Indentasi kiri-kanan supaya garisnya pendek (tidak selebar sel)
        pPr_ind = OxmlElement("w:ind")
        pPr_ind.set(qn("w:left"), "1400")
        pPr_ind.set(qn("w:right"), "1400")
        pPr.append(pPr_ind)

    # ------------------------------------------------------------------
    # QR CODE VERIFIKASI
    # ------------------------------------------------------------------
    def _tambah_qrcode(self):
        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        tmp.close()
        self._temp_files.append(tmp.name)

        data_verifikasi = (
            f"BUMDes:{self.info.get('nama_bumdes', '-')}|"
            f"Nomor:{self.info.get('nomor_surat', '-')}|"
            f"Periode:{self.info.get('periode', '-')}|"
            f"Cetak:{self.laporan.tanggal_cetak()}"
        )
        buat_qrcode(data_verifikasi, tmp.name)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("Kode verifikasi dokumen:")
        run.font.size = Pt(8)
        run.italic = True

        p_img = self.doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_img.add_run().add_picture(tmp.name, width=Cm(2.2))

    def _bersihkan_file_sementara(self):
        for f in self._temp_files:
            try:
                os.remove(f)
            except OSError:
                pass
